"""
tests/test_sii_defense_generator.py
===================================
Suite de pruebas unitarias para el Generador de Escritos de Descargo SII:
1. Validación de Márgenes y Estándar Formal de Documento Word (.docx).
2. Verificación de Citas Normativas (D.L. 825, D.L. 830, Circular 50).
3. Verificación de Identificación Territorial en Punta Arenas.
4. Cobertura de las 4 Causas Tributarias.
"""

import io
import pytest
import docx
from docx.shared import Cm
from core.sii.legal_defense_builder import SIILegalDefenseBuilder, format_monto_clp

MOCK_DATA_BASE = {
    "razon_social": "Gastronomía Magallánica SpA",
    "rut_empresa": "77.199.932-8",
    "rep_legal": "Regina Belén Andrade",
    "rep_rut": "18.902.386-3",
    "domicilio": "Calle Bories N° 450, Punta Arenas",
    "giro": "Restaurante y Servicios Gastronómicos",
    "periodos": ["2025-01", "2025-02", "2025-03"],
    "iva_declarado": 4500000.0,
    "numero_citacion": "Citación N° 402/2026",
    "argumentos_adicionales": "Se acompañan comprobantes de venta con tarjeta de débito y crédito Transbank."
}

def get_text_from_docx_stream(stream: io.BytesIO) -> str:
    doc = docx.Document(stream)
    return "\n".join(p.text for p in doc.paragraphs)

def test_sii_defense_doc_margins_and_format():
    stream = SIILegalDefenseBuilder.build_docx(MOCK_DATA_BASE)
    doc = docx.Document(stream)
    
    # Verificar márgenes reglamentarios (3.0 cm sup, 3.5 cm izq)
    sec = doc.sections[0]
    assert round(sec.top_margin.cm, 1) == 3.0
    assert round(sec.left_margin.cm, 1) == 3.5
    assert round(sec.bottom_margin.cm, 1) == 2.5
    assert round(sec.right_margin.cm, 1) == 2.5

def test_sii_defense_boletas_vs_facturas():
    data = {**MOCK_DATA_BASE, "doc_type": "boletas_vs_facturas"}
    stream = SIILegalDefenseBuilder.build_docx(data)
    text = get_text_from_docx_stream(stream)
    
    assert "Director Regional" in text
    assert "XII Región de Magallanes" in text
    assert "Unidad Punta Arenas" in text
    assert "GASTRONOMÍA MAGALLÁNICA SPA" in text.upper()
    assert "77.199.932-8" in text
    assert "artículo 53 del D.L. N° 825" in text
    assert "$4.500.000" in text
    assert "POR TANTO" in text
    assert "REGINA BELÉN ANDRADE" in text.upper()

def test_sii_defense_citacion_art_63():
    data = {**MOCK_DATA_BASE, "doc_type": "citacion_art_63"}
    stream = SIILegalDefenseBuilder.build_docx(data)
    text = get_text_from_docx_stream(stream)
    
    assert "CITACIÓN N° Citación N° 402/2026 (ART. 63 CÓDIGO TRIBUTARIO)" in text
    assert "Registro de Compras y Ventas (RCV)" in text

def test_sii_defense_rectificatoria_f29():
    data = {**MOCK_DATA_BASE, "doc_type": "rectificatoria_f29"}
    stream = SIILegalDefenseBuilder.build_docx(data)
    text = get_text_from_docx_stream(stream)
    
    assert "artículo 127 del D.L. N° 830 (Código Tributario)" in text
    assert "rectificatoria voluntaria" in text.lower()
    assert "error de hecho" in text.lower()

def test_sii_defense_condonacion_multas():
    data = {**MOCK_DATA_BASE, "doc_type": "condonacion_multas"}
    stream = SIILegalDefenseBuilder.build_docx(data)
    text = get_text_from_docx_stream(stream)
    
    assert "Circular N° 50 del Servicio de Impuestos Internos" in text
    assert "condonación del máximo porcentaje legal" in text

def test_format_monto_clp():
    assert format_monto_clp(1250000) == "$1.250.000"
    assert format_monto_clp("345000") == "$345.000"
    assert format_monto_clp(0) == "$0"
    assert format_monto_clp(None) == "$0"
