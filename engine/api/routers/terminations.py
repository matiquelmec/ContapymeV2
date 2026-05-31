from datetime import date, datetime, timedelta
import calendar
from typing import Optional, List, Dict, Any
from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from core.database import get_supabase
from core.auth import verify_token, verify_org_role
import os
import tempfile
import uuid
import base64
import io
import requests

router = APIRouter()

class TerminationRequest(BaseModel):
    employee_id: str
    organization_id: str
    fecha_termino: date
    causal_despido: str 
    aviso_previo: bool = True
    dias_vacaciones_tomados: float = 0
    pending_overtime_amount: int = 0
    other_bonuses: int = 0
    asignacion_colacion: int = 0
    asignacion_movilizacion: int = 0
    viaticos: int = 0
    prestamo_ccaf: int = 0
    anticipo_sueldo: int = 0
    banco_transferencia: Optional[str] = ""
    tipo_cuenta: Optional[str] = ""
    cuenta_transferencia: Optional[str] = ""

def calculate_years_of_service(start_date: date, end_date: date):
    delta = end_date - start_date
    total_days = delta.days + 1
    
    years = total_days // 365
    remaining_days = total_days % 365
    months = remaining_days // 30
    
    severance_years = years
    if years >= 1 and (months > 6 or (months == 6 and (remaining_days % 30) > 0)):
        severance_years += 1
        
    return {
        "years": years,
        "months": months,
        "total_days": total_days,
        "severance_years": min(severance_years, 11)
    }

def calculate_proportional_holidays_precise(start_date: date, end_date: date):
    # Fórmula estricta de vacaciones proporcionales DT (Días corridos * (1.25/30))
    delta = end_date - start_date
    total_days = delta.days + 1
    proportional_days = (total_days / 30.0) * 1.25
    return round(proportional_days, 2)

@router.post("/calculate")
async def calculate_termination(
    req: TerminationRequest,
    current_user: dict = Depends(verify_token)
):
    await verify_org_role(req.organization_id, auth=current_user)
    db = get_supabase()
    try:
        emp_res = db.table("employees").select("*").eq("id", req.employee_id).single().execute()
        emp = emp_res.data
        if not emp:
            raise HTTPException(status_code=404, detail="Empleado no encontrado")

        fecha_ingreso = date.fromisoformat(emp["fecha_ingreso"])
        
        # Base de Cálculo Finiquito (Art. 172 CT):
        # Sueldo Base + Gratificación + Colación + Movilización (Valores Fijos y Constantes)
        base_calculo = emp.get("sueldo_base", 0) + emp.get("asignacion_colacion", 0) + emp.get("asignacion_movilizacion", 0)
        if emp.get("gratificacion_legal", True):
            # Asumimos el tope de ~200.000 como referencia base si no hay DB
            base_calculo += min(int(emp.get("sueldo_base", 0) * 0.25), 209000)
            
        # El sueldo diario para vacaciones usa solo Sueldo Base (Art. 71)
        sueldo_base = emp.get("sueldo_base", 0)
        
        uf_value = 37000 
        try:
            ind_res = db.table("economic_indicators").select("valor").eq("codigo", "uf").execute()
            if ind_res.data:
                uf_value = float(ind_res.data[0]["valor"])
        except:
            pass
        
        requires_severance = False
        requires_notice = False
        severance_calculation_type = None
        causal_desc = req.causal_despido
 
        try:
            cause_res = db.table("termination_causes").select("*").or_(f"article_code.eq.{req.causal_despido},article_name.ilike.%{req.causal_despido}%").execute()
            if cause_res.data:
                cause = cause_res.data[0]
                causal_desc = cause["article_name"]
                requires_severance = cause["requires_severance"]
                requires_notice = cause["requires_notice"]
                severance_calculation_type = cause["severance_calculation_type"]
            else:
                if "161" in req.causal_despido or "necesidades" in req.causal_despido.lower():
                    requires_severance = True
                    requires_notice = True
                    severance_calculation_type = "years_service"
        except:
            if "161" in req.causal_despido or "necesidades" in req.causal_despido.lower():
                requires_severance = True
                requires_notice = True
                severance_calculation_type = "years_service"
 
        time_stats = calculate_years_of_service(fecha_ingreso, req.fecha_termino)
        
        last_day_of_month = calendar.monthrange(req.fecha_termino.year, req.fecha_termino.month)[1]
        worked_days_last_month = req.fecha_termino.day
        pending_salary_amount = int((sueldo_base / 30) * worked_days_last_month)
        
        proportional_days = calculate_proportional_holidays_precise(fecha_ingreso, req.fecha_termino)
        pending_vacation_days = max(0.0, float(proportional_days) - float(req.dias_vacaciones_tomados))
        vacation_daily_rate = int(sueldo_base / 30)
        monto_vacaciones = int(pending_vacation_days * vacation_daily_rate)
        
        monto_anos_servicio = 0
        if requires_severance and severance_calculation_type == "years_service":
            # TOPE LEGAL ART 172: El tope mensual para indemnizaciones es 90 UF
            tope_90_uf = int(uf_value * 90)
            base_indemnizacion = min(base_calculo, tope_90_uf)
            monto_anos_servicio = base_indemnizacion * time_stats["severance_years"]
 
        monto_mes_aviso = 0
        if requires_notice and not req.aviso_previo:
            tope_90_uf_aviso = int(uf_value * 90)
            base_aviso = min(base_calculo, tope_90_uf_aviso)
            monto_mes_aviso = base_aviso
 
        # Total bruto de compensaciones (haberes de finiquito)
        total_haberes_finiquito = (
            pending_salary_amount + 
            monto_vacaciones + 
            monto_anos_servicio + 
            monto_mes_aviso + 
            req.pending_overtime_amount + 
            req.other_bonuses +
            req.asignacion_colacion +
            req.asignacion_movilizacion +
            req.viaticos
        )
        
        # Descuentos de finiquito
        total_descuentos_finiquito = req.prestamo_ccaf + req.anticipo_sueldo
        
        # Total líquido final de finiquito
        total_liquido_finiquito = total_haberes_finiquito - total_descuentos_finiquito
        
        termination_data = {
            "organization_id": req.organization_id,
            "employee_id": req.employee_id,
            "fecha_inicio": fecha_ingreso.isoformat(),
            "fecha_termino": req.fecha_termino.isoformat(),
            "causal_despido": causal_desc,
            "vacaciones_pendientes_dias": float(pending_vacation_days),
            "monto_vacaciones": int(monto_vacaciones),
            "monto_indemnizacion_anos": int(monto_anos_servicio),
            "monto_mes_aviso": int(monto_mes_aviso),
            "total_finiquito": int(total_liquido_finiquito),
            "status": "borrador",
            "worked_days_last_month": worked_days_last_month,
            "pending_salary_amount": pending_salary_amount,
            "vacation_daily_rate": vacation_daily_rate,
            "proportional_vacation_days": float(proportional_days),
            "proportional_vacation_amount": int(monto_vacaciones),
            "severance_years_service": float(time_stats["severance_years"]),
            "severance_monthly_salary": int(min(base_calculo, int(uf_value * 90))),
            "notice_indemnification_amount": int(monto_mes_aviso),
            "pending_overtime_amount": req.pending_overtime_amount,
            "other_bonuses_amount": req.other_bonuses,
            "asignacion_colacion": req.asignacion_colacion,
            "asignacion_movilizacion": req.asignacion_movilizacion,
            "viaticos": req.viaticos,
            "prestamo_ccaf": req.prestamo_ccaf,
            "anticipo_sueldo": req.anticipo_sueldo,
            "updated_at": datetime.now().isoformat()
        }
 
        # Actualizar datos bancarios directamente en la ficha del empleado (Employees) para normalización
        db.table("employees").update({
            "banco_transferencia": req.banco_transferencia,
            "tipo_cuenta": req.tipo_cuenta,
            "cuenta_transferencia": req.cuenta_transferencia
        }).eq("id", req.employee_id).execute()
 
        exist = db.table("employee_terminations").select("id").eq("employee_id", req.employee_id).execute()
        
        if exist.data:
            term_id = exist.data[0]["id"]
            db.table("employee_terminations").update(termination_data).eq("id", term_id).execute()
            termination_data["id"] = term_id
        else:
            ins_res = db.table("employee_terminations").insert(termination_data).execute()
            termination_data["id"] = ins_res.data[0]["id"]
 
        return {
            "success": True,
            "data": termination_data,
            "calculation_details": {
                "years_of_service": time_stats["years"],
                "months_of_service": time_stats["months"],
                "uf_used": uf_value,
                "severance_years_applied": time_stats["severance_years"]
            }
        }
    except Exception as e:
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Error en el cálculo de finiquito: {str(e)}")

@router.get("/causes")
async def get_termination_causes(current_user: dict = Depends(verify_token)):
    db = get_supabase()
    try:
        res = db.table("termination_causes").select("*").order("article_code").execute()
        if not res.data:
            return {
                "success": True, 
                "data": [
                    {"article_code": "161-1", "article_name": "Art. 161 N°1 - Necesidades de la empresa", "requires_notice": True, "requires_severance": True},
                    {"article_code": "159-1", "article_name": "Art. 159 N°1 - Mutuo acuerdo", "requires_notice": False, "requires_severance": False},
                    {"article_code": "159-2", "article_name": "Art. 159 N°2 - Renuncia voluntaria", "requires_notice": False, "requires_severance": False},
                    {"article_code": "159-4", "article_name": "Art. 159 N°4 - Vencimiento del plazo", "requires_notice": False, "requires_severance": False},
                    {"article_code": "159-5", "article_name": "Art. 159 N°5 - Conclusión de obra", "requires_notice": False, "requires_severance": False}
                ]
            }
        return {"success": True, "data": res.data}
    except:
        return {
            "success": True, 
            "data": [
                {"article_code": "161-1", "article_name": "Art. 161 N°1 - Necesidades de la empresa", "requires_notice": True, "requires_severance": True},
                {"article_code": "159-1", "article_name": "Art. 159 N°1 - Mutuo acuerdo", "requires_notice": False, "requires_severance": False}
            ]
        }

def format_date_spanish(dt: date | datetime) -> str:
    meses = {
        1: "Enero", 2: "Febrero", 3: "Marzo", 4: "Abril",
        5: "Mayo", 6: "Junio", 7: "Julio", 8: "Agosto",
        9: "Septiembre", 10: "Octubre", 11: "Noviembre", 12: "Diciembre"
    }
    return f"{dt.day} de {meses[dt.month]} de {dt.year}"

@router.get("/{termination_id}/document/{doc_type}")
async def generate_document_text(
    termination_id: str, 
    doc_type: str,
    current_user: dict = Depends(verify_token)
):
    db = get_supabase()
    try:
        term_res = db.table("employee_terminations").select("*, employees(*)").eq("id", termination_id).single().execute()
        term = term_res.data
        if not term:
            raise HTTPException(status_code=404, detail="Finiquito no encontrado")
        
        await verify_org_role(term["organization_id"], auth=current_user)
        
        emp = term.get("employees", {})
        if not emp:
            emp_res = db.table("employees").select("*").eq("id", term["employee_id"]).single().execute()
            emp = emp_res.data or {}

        org_id = term["organization_id"]
        org_res = db.table("organizations").select("*").eq("id", org_id).single().execute()
        org = org_res.data or {}
        
        config_res = db.table("organization_payroll_settings").select("*").eq("organization_id", org_id).maybe_single().execute()
        config = config_res.data or {}
        rep_nombre = config.get("rep_legal_nombre", "________________")
        rep_rut = config.get("rep_legal_rut", "________________")
        
        hoy = format_date_spanish(date.today())
        
        # Robust date formatting helper
        def get_fmt_date(val):
            if not val: return "---"
            if isinstance(val, str):
                try:
                    f_iso = val.split('T')[0]
                    return format_date_spanish(datetime.strptime(f_iso, "%Y-%m-%d"))
                except: return str(val)
            return format_date_spanish(val)

        fecha_term = get_fmt_date(term.get("fecha_termino"))
        fecha_inicio = get_fmt_date(term.get("fecha_inicio"))
        
        if doc_type == "carta":
            return {
                "success": True,
                "data": {
                    "title": f"CARTA DE AVISO - {emp.get('nombres', '')} {emp.get('apellido_paterno', '')}",
                    "content": (
                        f"{org.get('comuna', 'Punta Arenas')}, {hoy}\n\n"
                        f"Señor(a)\n{emp.get('nombres', '')} {emp.get('apellido_paterno', '')} {emp.get('apellido_materno', '')}\n"
                        f"RUT: {emp.get('rut', '---')}\nPresente:\n\n"
                        f"Ref: Aviso de Término de Contrato de Trabajo.\n\n"
                        f"Por medio de la presente, comunicamos a Ud. que la empresa {org.get('nombre', 'LA EMPRESA')} "
                        f"ha decidido poner término a su contrato de trabajo, en virtud de lo dispuesto en el "
                        f"{term.get('causal_despido', 'Artículo correspondiente')}, del Código del Trabajo.\n\n"
                        f"El término de su contrato será efectivo a partir del día {fecha_term}.\n\n"
                        f"Asimismo, ponemos en su conocimiento que sus cotizaciones previsionales y de salud se encuentran al día, "
                        f"según consta en certificados adjuntos.\n\n"
                        f"Finalmente, agradecemos su desempeño y compromiso durante el período en que le correspondió cumplir funciones en nuestra institución.\n\n"
                        f"Sin otro particular, saluda atentamente a Ud.\n\n"
                        f"__________________________\nFIRMA EMPLEADOR\n{rep_nombre}\nRUT: {rep_rut}\np.p. {org.get('nombre', 'CONTAPYME V2')}"
                    )
                }
            }
        elif doc_type == "finiquito":
            transfer_text = ""
            banco = emp.get("banco_transferencia") or term.get("banco_transferencia")
            tipo_cta = emp.get("tipo_cuenta") or term.get("tipo_cuenta")
            nro_cta = emp.get("cuenta_transferencia") or term.get("cuenta_transferencia")
            
            if nro_cta:
                transfer_text = f"\n\nEl pago se realiza mediante transferencia electrónica a la cuenta {tipo_cta or 'corriente'} N° {nro_cta} del {banco or 'Banco Estado'}."
            
            return {
                "success": True,
                "data": {
                    "title": f"ACTA DE FINIQUITO - {emp.get('nombres', '')} {emp.get('apellido_paterno', '')}",
                    "content": (
                        f"En {org.get('comuna', 'Punta Arenas')}, a {hoy}, entre la empresa {org.get('nombre', 'LA EMPRESA')}, "
                        f"RUT {org.get('rut_empresa', '---')}, representada por su Gerente General, y el trabajador "
                        f"don (ña) {emp.get('nombres', '')} {emp.get('apellido_paterno', '')}, RUT {emp.get('rut', '---')}, "
                        f"se ha convenido el siguiente finiquito:\n\n"
                        f"PRIMERO: El trabajador prestó servicios desde el {fecha_inicio} hasta el {fecha_term}, "
                        f"fecha en que el contrato ha terminado por la causal: {term.get('causal_despido', '---')}.\n\n"
                        f"SEGUNDO: Por medio de este acto, el empleador paga al trabajador las siguientes sumas:\n"
                        f"HABERES:\n"
                        f"- Haberes pendientes y días trabajados: ${term.get('pending_salary_amount', 0):,}\n"
                        f"- Vacaciones proporcionales / pendientes: ${term.get('monto_vacaciones', 0):,}\n"
                        f"- Indemnización años de servicio: ${term.get('monto_indemnizacion_anos', 0):,}\n"
                        f"- Mes de aviso / Sustitutiva: ${term.get('monto_mes_aviso', 0):,}\n"
                        f"- Otros bonos y horas extras: ${term.get('pending_overtime_amount', 0) + term.get('other_bonuses_amount', 0):,}\n"
                        f"- Asignación de Colación: ${term.get('asignacion_colacion', 0):,}\n"
                        f"- Asignación de Movilización: ${term.get('asignacion_movilizacion', 0):,}\n"
                        f"- Viáticos / Otros asignaciones: ${term.get('viaticos', 0):,}\n"
                        f"DESCUENTOS:\n"
                        f"- Préstamo / Crédito CCAF: ${term.get('prestamo_ccaf', 0):,}\n"
                        f"- Anticipos de Sueldo: ${term.get('anticipo_sueldo', 0):,}\n\n"
                        f"TOTAL LÍQUIDO A PAGAR: ${term.get('total_finiquito', 0):,}{transfer_text}\n\n"
                        f"TERCERO: El trabajador declara recibir en este acto, a su entera satisfacción, la suma total indicada, "
                        f"no teniendo reclamo alguno que formular en contra de su empleador derivado de la relación laboral que los unió.\n\n"
                        f"CUARTO: Las partes otorgan el más amplio, completo y recíproco finiquito.\n\n"
                        f"__________________________\n"
                        f"FIRMA EMPLEADOR\n"
                        f"{rep_nombre}\n"
                        f"RUT: {rep_rut}\n"
                        f"p.p. {org.get('nombre', '---')}\n\n"
                        f"__________________________\n"
                        f"FIRMA TRABAJADOR\n"
                        f"{emp.get('nombres', '')} {emp.get('apellido_paterno', '')}\n"
                        f"RUT: {emp.get('rut', '---')}\n"
                    )
                }
            }
        return {"success": False, "error": "Tipo de documento no válido"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/{termination_id}/download/{doc_type}")
async def download_docx(
    termination_id: str, 
    doc_type: str,
    current_user: dict = Depends(verify_token)
):
    db = get_supabase()
    try:
        # Reutilizar lógica de obtención de datos
        term_res = db.table("employee_terminations").select("*, employees(*)").eq("id", termination_id).single().execute()
        term = term_res.data
        if not term:
            raise HTTPException(status_code=404, detail="Finiquito no encontrado")
        
        await verify_org_role(term["organization_id"], auth=current_user)
        
        emp = term.get("employees", {})
        org_id = term["organization_id"]

        org_res = db.table("organizations").select("*").eq("id", org_id).single().execute()
        org = org_res.data or {}

        config_res = db.table("organization_payroll_settings").select("*").eq("organization_id", org_id).maybe_single().execute()
        config = config_res.data or {}
        rep_nombre = config.get("rep_legal_nombre", "________________")
        rep_rut = config.get("rep_legal_rut", "________________")
        
        # Obtener textos (reutilizando la lógica anterior o generando nuevos)
        hoy = format_date_spanish(date.today())
        if isinstance(term["fecha_termino"], str):
            f_iso = term["fecha_termino"].split('T')[0]
            dt_term = datetime.strptime(f_iso, "%Y-%m-%d")
            fecha_term = format_date_spanish(dt_term)
        else:
            fecha_term = format_date_spanish(term["fecha_termino"])

        sig_base64 = term.get("signature_base64")

        # Crear documento Word
        doc = Document()
        
        # Estilo base
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

        if doc_type == "carta":
            # Encabezado Ciudad y Fecha
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(f"{org.get('comuna', 'Punta Arenas')}, {hoy}").bold = True
            
            doc.add_paragraph("\nSeñor(a)")
            doc.add_paragraph(f"{emp['nombres']} {emp['apellido_paterno']} {emp['apellido_materno'] or ''}").bold = True
            doc.add_paragraph(f"RUT: {emp['rut']}")
            doc.add_paragraph("Presente:\n")
            
            ref = doc.add_paragraph()
            ref.add_run("Ref: Aviso de Término de Contrato de Trabajo.").bold = True
            ref.add_run("\n")

            body = doc.add_paragraph()
            body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            body.add_run(f"Por medio de la presente, comunicamos a Ud. que la empresa {org.get('nombre', 'LA EMPRESA')} ha decidido poner término a su contrato de trabajo, en virtud de lo dispuesto en el ")
            body.add_run(f"{term['causal_despido']}").bold = True
            body.add_run(", del Código del Trabajo.")
            
            body2 = doc.add_paragraph()
            body2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            body2.add_run(f"El término de su contrato será efectivo a partir del día {fecha_term}.")
            
            doc.add_paragraph("\nAsimismo, ponemos en su conocimiento que sus cotizaciones previsionales y de salud se encuentran al día, según consta en certificados adjuntos.")
            doc.add_paragraph("Finalmente, agradecemos su desempeño y compromiso durante el período en que le correspondió cumplir funciones en nuestra institución.")
            
            doc.add_paragraph("\nSin otro particular, saluda atentamente a Ud.\n\n")
            
            # Tabla de firmas para Carta
            table = doc.add_table(rows=1, cols=1) # Solo empleador en la carta
            table.autofit = True
            cell = table.rows[0].cells[0]
            
            f1 = cell.paragraphs[0]
            f1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            f1.add_run("__________________________\nFIRMA EMPLEADOR\n").bold = True
            f1.add_run(f"{rep_nombre}\n")
            f1.add_run(f"RUT: {rep_rut}\n")
            f1.add_run(f"p.p. {org.get('nombre', 'CONTAPYMEPUQ')}").italic = True

            # --- INCORPORAR FIRMA DIGITAL SI EXISTE ---
            if sig_base64:
                try:
                    if "," in sig_base64:
                        sig_base64 = sig_base64.split(",")[1]
                    sig_data = base64.b64decode(sig_base64)
                    sig_stream = io.BytesIO(sig_data)
                    
                    p_sig = doc.add_paragraph()
                    p_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Firma arriba
                    run_img = p_sig.add_run()
                    run_img.add_picture(sig_stream, width=Inches(1.5))
                    
                    # --- GENERAR QR DE VERIFICACIÓN ---
                    try:
                        verify_url = f"https://contapymepuq.cl/verify/{emp.get('rut', 'ID')}"
                        qr_api_url = f"https://api.qrserver.com/v1/create-qr-code/?size=150x150&data={verify_url}"
                        qr_response = requests.get(qr_api_url)
                        if qr_response.status_code == 200:
                            qr_stream = io.BytesIO(qr_response.content)
                            p_sig.add_run("   ") # Espacio
                            p_sig.add_run().add_picture(qr_stream, width=Inches(0.6))
                    except Exception as qr_err:
                        print(f"⚠️ Error al generar QR: {qr_err}")

                    # Línea y Sello debajo
                    p_sig.add_run("\n__________________________\n").bold = True
                    p_sig.add_run("VERIFICACIÓN DE IDENTIDAD DIGITAL\n").bold = True
                    p_sig.add_run("CONTAPYMEPUQ - SELLO DE TIEMPO\n").italic = True
                except Exception as e:
                    print(f"Error embedding signature in carta: {e}")

            # Pie de página Premium
            doc.add_paragraph("\n\n")
            footer_p = doc.add_paragraph()
            footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = footer_p.add_run(f"Documento generado por el sistema de gestión laboral Contapymepuq\n{org.get('nombre', 'Empresa')} — {org.get('comuna', 'Punta Arenas')}")
            run.font.size = Pt(8)
            run.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128) # Gris profesional

        elif doc_type == "finiquito":
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.add_run("ACTA DE FINIQUITO DE CONTRATO DE TRABAJO").bold = True
            
            doc.add_paragraph(f"\nEn {org.get('comuna', 'Punta Arenas')}, a {hoy}, entre la empresa {org.get('nombre', 'LA EMPRESA')}, RUT {org.get('rut_empresa', '---')}, representada por su Gerente General, y el trabajador don (ña) {emp['nombres']} {emp['apellido_paterno']}, RUT {emp['rut']}, se ha convenido el siguiente finiquito:\n")
            
            p1 = doc.add_paragraph()
            p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p1.add_run("PRIMERO: ").bold = True
            p1.add_run(f"El trabajador prestó servicios desde el {term['fecha_inicio']} hasta el {term['fecha_termino']}, fecha en que el contrato ha terminado por la causal: {term['causal_despido']}.")
            
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p2.add_run("SEGUNDO: ").bold = True
            p2.add_run("Por medio de este acto, el empleador paga al trabajador las siguientes sumas:\n")
            
            items = [
                f"- Haberes pendientes y días trabajados: ${term.get('pending_salary_amount', 0):,}",
                f"- Vacaciones proporcionales / pendientes: ${term['monto_vacaciones']:,}",
                f"- Indemnización años de servicio: ${term['monto_indemnizacion_anos']:,}",
                f"- Mes de aviso / Sustitutiva: ${term['monto_mes_aviso']:,}",
                f"- Otros bonos y horas extras: ${term.get('pending_overtime_amount', 0) + term.get('other_bonuses_amount', 0):,}"
            ]
            for item in items:
                li = doc.add_paragraph(item)
                li.paragraph_format.left_indent = Inches(0.5)
            
            total = doc.add_paragraph()
            total.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            total.add_run(f"\nTOTAL BRUTO A PAGAR: ${term['total_finiquito']:,}").bold = True
            
            p3 = doc.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p3.add_run("\nTERCERO: ").bold = True
            p3.add_run("El trabajador declara recibir en este acto, a su entera satisfacción, la suma total indicada, no teniendo reclamo alguno que formular en contra de su empleador derivado de la relación laboral que los unió.")
            
            p4 = doc.add_paragraph()
            p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p4.add_run("CUARTO: ").bold = True
            p4.add_run("Las partes otorgan el más amplio, completo y recíproco finiquito.\n\n")

            # Tabla de firmas
            table = doc.add_table(rows=1, cols=2)
            table.autofit = True
            cells = table.rows[0].cells
            
            f1 = cells[0].paragraphs[0]
            f1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            f1.add_run("__________________________\nFIRMA EMPLEADOR\n").bold = True
            f1.add_run(f"{rep_nombre}\n")
            f1.add_run(f"RUT: {rep_rut}\n")
            f1.add_run(f"p.p. {org.get('nombre', 'CONTAPYME V2')}").italic = True
            
            f2 = cells[1].paragraphs[0]
            f2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            f2.add_run("__________________________\nFIRMA TRABAJADOR\n").bold = True
            f2.add_run(f"{emp.get('nombres', '')} {emp.get('apellido_paterno', '')}\n")
            f2.add_run(f"RUT: {emp.get('rut', '---')}")

            # --- INCORPORAR FIRMA DIGITAL EN FINIQUITO ---
            if sig_base64:
                try:
                    if "," in sig_base64:
                        sig_base64 = sig_base64.split(",")[1]
                    sig_data = base64.b64decode(sig_base64)
                    sig_stream = io.BytesIO(sig_data)
                    
                    p_sig = doc.add_paragraph()
                    p_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_sig.add_run("\nPROTOCOLO DE FIRMA DIGITAL CERTIFICADA\n").bold = True
                    run_img = p_sig.add_run()
                    run_img.add_picture(sig_stream, width=Inches(2.0))
                except Exception as e:
                    print(f"Error embedding signature in finiquito: {e}")

            # Pie de página Premium
            doc.add_paragraph("\n\n")
            footer_p = doc.add_paragraph()
            footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            full_address_fin = (lambda d, c: d if c.lower() in d.lower() else f"{d}, {c}")(
                org.get('direccion', ''), 
                org.get('comuna', 'Punta Arenas')
            )
            run = footer_p.add_run(f"Documento generado por el sistema de gestión laboral Contapymepuq\n{org.get('nombre', 'Empresa')} — {full_address_fin}")
            run.font.size = Pt(8)
            run.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128) # Gris profesional

        # Guardar en archivo temporal
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        
        filename = f"{doc_type}_{emp['apellido_paterno']}_{termination_id[:8]}.docx"
        
        return FileResponse(
            path=tmp.name,
            filename=filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except Exception as e:
        print(f"[DOCX ERROR] {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/{termination_id}/export-dt-csv")
async def export_dt_csv(
    termination_id: str,
    current_user: dict = Depends(verify_token)
):
    """
    Genera el CSV de carga masiva de Finiquito Electrónico de la Dirección del Trabajo (48 columnas).
    Basado en el instructivo oficial DT versión 3.0.
    """
    db = get_supabase()
    try:
        term_res = db.table("employee_terminations").select("*, employees(*)").eq("id", termination_id).single().execute()
        term = term_res.data
        if not term:
            raise HTTPException(status_code=404, detail="Finiquito no encontrado")
        
        await verify_org_role(term["organization_id"], auth=current_user)
        
        emp = term.get("employees", {})
        org_id = term["organization_id"]

        org_res = db.table("organizations").select("*").eq("id", org_id).single().execute()
        org = org_res.data or {}

        # Mapeos de códigos oficiales de la Dirección del Trabajo
        CAUSALES_DT = {
            "Art. 159 N°1": 3, "MUTUO ACUERDO": 3,
            "Art. 159 N°2": 4, "RENUNCIA": 4,
            "Art. 159 N°4": 6, "VENCIMIENTO DEL PLAZO": 6,
            "Art. 159 N°5": 7, "CONCLUSION DE OBRA": 7,
            "Art. 161 N°1": 18, "NECESIDADES DE LA EMPRESA": 18
        }
        
        causal_raw = (term.get("causal_despido") or "").upper()
        causal_id = 18 # Fallback: Necesidades de la empresa
        for k, v in CAUSALES_DT.items():
            if k in causal_raw:
                causal_id = v
                break

        # Regiones y Comunas
        # Mapeo simple de código de región para la DT
        region_str = (org.get("region") or "XII").upper()
        region_id = 12 # Magallanes default
        if "METROPOLITANA" in region_str or "RM" in region_str: region_id = 13
        elif "AYSEN" in region_str: region_id = 11

        comuna_id = 12101 # Punta Arenas default

        # 48 Columnas requeridas por la DT
        headers = [
            "RutEmpresa", "RutTrabajador", "FechaInicioContrato", "FechaTerminoContrato",
            "DeclaraNotificacionRetencionAlimento", "CausalFiniquitoId", "Funciones",
            "RegionTrabajoId", "ComunaId", "LugarPrestacionServiciosDireccion",
            "CantidadDiasVacaciones", "IndemnizacionFeriado", "IndemnizacionAvisoPrevio",
            "IndemnizacionServicio", "IndemnizacionOtras", "IndemnizacionArticulo163",
            "remuneracionPendiente", "Gratificaciones", "Bonos", "HorasExtraordinarias",
            "Aguinaldo", "SemanaCorrida", "ComisionOParticipacion", "Movilizacion",
            "Colacion", "PerdidaCaja", "DesgasteHerramientas", "Viaticos",
            "AsignacionesFamiliares", "DescuentoSeguridadSocial", "DescuentoImpuestos",
            "DescuentoAfc", "DescuentoAnticipado", "DescuentoIndemnizacion",
            "DescuentoPension", "DescuentoCajaCompensacion", "PrestamoAdeudado",
            "AnticipoSueldo", "VacacionesAnticipadas", "Email", "CodigoComunaPersonal",
            "CallePersonal", "NumeroPersonal", "DepartamentoBlockPersonal", "Telefono",
            "CuentaTransferencia", "BancoId", "TipoCuentaId"
        ]

        def fmt_rut(r: str) -> str:
            # Formato sin puntos pero con guión
            limpio = r.replace(".", "").upper()
            if "-" not in limpio and len(limpio) > 1:
                return f"{limpio[:-1]}-{limpio[-1]}"
            return limpio

        def fmt_date(d_str: str) -> str:
            if not d_str: return ""
            # YYYY-MM-DD -> DD-MM-YYYY
            parts = d_str.split("T")[0].split("-")
            if len(parts) == 3:
                return f"{parts[2]}-{parts[1]}-{parts[0]}"
            return d_str

        # Poblar valores
        data = {
            "RutEmpresa": fmt_rut(org.get("rut_empresa", "")),
            "RutTrabajador": fmt_rut(emp.get("rut", "")),
            "FechaInicioContrato": fmt_date(term.get("fecha_inicio")),
            "FechaTerminoContrato": fmt_date(term.get("fecha_termino")),
            "DeclaraNotificacionRetencionAlimento": "No",
            "CausalFiniquitoId": str(causal_id),
            "Funciones": (emp.get("cargo") or "OPERARIO")[:100],
            "RegionTrabajoId": str(region_id),
            "ComunaId": str(comuna_id),
            "LugarPrestacionServiciosDireccion": (org.get("direccion") or "Punta Arenas")[:100],
            "CantidadDiasVacaciones": str(int(term.get("vacaciones_pendientes_dias", 0))),
            "IndemnizacionFeriado": str(term.get("monto_vacaciones", 0)),
            "IndemnizacionAvisoPrevio": str(term.get("monto_mes_aviso", 0)),
            "IndemnizacionServicio": str(term.get("monto_indemnizacion_anos", 0)),
            "IndemnizacionOtras": "0",
            "IndemnizacionArticulo163": "0",
            "remuneracionPendiente": str(term.get("pending_salary_amount", 0)),
            "Gratificaciones": "0",
            "Bonos": str(term.get("other_bonuses_amount", 0)),
            "HorasExtraordinarias": str(term.get("pending_overtime_amount", 0)),
            "Aguinaldo": "0",
            "SemanaCorrida": "0",
            "ComisionOParticipacion": "0",
            "Movilizacion": str(term.get("asignacion_movilizacion", 0)),
            "Colacion": str(term.get("asignacion_colacion", 0)),
            "PerdidaCaja": "0",
            "DesgasteHerramientas": "0",
            "Viaticos": str(term.get("viaticos", 0)),
            "AsignacionesFamiliares": "0",
            "DescuentoSeguridadSocial": "0",
            "DescuentoImpuestos": "0",
            "DescuentoAfc": "0",
            "DescuentoAnticipado": "0",
            "DescuentoIndemnizacion": "0",
            "DescuentoPension": "0",
            "DescuentoCajaCompensacion": str(term.get("prestamo_ccaf", 0)),
            "PrestamoAdeudado": "0",
            "AnticipoSueldo": str(term.get("anticipo_sueldo", 0)),
            "VacacionesAnticipadas": "0",
            "Email": (emp.get("email") or "")[:100],
            "CodigoComunaPersonal": str(comuna_id),
            "CallePersonal": (emp.get("address") or "Calle Principal")[:50],
            "NumeroPersonal": "100",
            "DepartamentoBlockPersonal": "",
            "Telefono": (emp.get("phone") or "999999999")[-9:],
            "CuentaTransferencia": emp.get("cuenta_banco", "") or term.get("cuenta_transferencia", ""),
            "BancoId": "3",  # Estado default
            "TipoCuentaId": "4"  # Rut default
        }

        # Generar CSV en memoria
        output = io.StringIO()
        import csv
        writer = csv.writer(output, delimiter=",", quoting=csv.QUOTE_MINIMAL)
        writer.writerow(headers)
        writer.writerow([data.get(h, "") for h in headers])
        
        content = output.getvalue()
        output.close()

        # Retornar descarga con UTF-8-sig para que Excel abra correctamente con acentos
        filename = f"FINIQUITO_DT_{emp.get('apellido_paterno', 'TRAB')}.csv"
        
        from fastapi.responses import StreamingResponse
        return StreamingResponse(
            iter([content.encode("utf-8-sig")]), 
            media_type="text/csv", 
            headers={'Content-Disposition': f'attachment; filename="{filename}"'}
        )
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))
