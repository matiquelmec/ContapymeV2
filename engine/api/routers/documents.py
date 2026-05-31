import os
import traceback
import base64
import io
import requests
from datetime import date, datetime
from io import BytesIO
from typing import Optional
from pydantic import BaseModel
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxtpl import DocxTemplate
from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import StreamingResponse

from core.database import get_supabase
from core.auth import verify_token, verify_org_role
from core.utils.shared_utils import clean_rut, format_clp, to_words, format_date_cl

router = APIRouter()

# Las plantillas vivirán en la carpeta engine/templates/
TEMPLATES_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "templates")

class GenerateDocRequest(BaseModel):
    employee_id: str
    type: str = "contrato"
    description: Optional[str] = ""
    signature_base64: Optional[str] = None

@router.post("/generate")
async def generate_document(
    req: GenerateDocRequest,
    current_user: dict = Depends(verify_token)
):
    """
    Genera un Contrato o Anexo de Trabajo inteligente.
    """
    employee_id = req.employee_id
    type = req.type
    description = req.description
    sig_raw = req.signature_base64
    
    db = get_supabase()

    try:
        # 1. Obtener datos del empleado y su empresa
        result = db.table("employees") \
            .select("*, organizations(*)") \
            .eq("id", employee_id) \
            .maybe_single() \
            .execute()

        emp = result.data
        if not emp:
            raise HTTPException(status_code=404, detail="Empleado no encontrado")
            
        org = emp.get('organizations', {})
        if isinstance(org, list) and len(org) > 0:
            org = org[0]
        
        org_id = org.get('id') or emp.get('organization_id')
        if org_id:
            await verify_org_role(org_id, auth=current_user)
        
        # 2. Cargar parámetros legales (Rep Legal)
        settings = {}
        if org_id:
            settings_res = db.table("organization_payroll_settings").select("*").eq("organization_id", org_id).maybe_single().execute()
            settings = settings_res.data or {}
        
        # 3. Lógica "Inteligente": Transformación de datos
        sueldo_base = emp.get('sueldo_base', 0)
        colacion = emp.get('asignacion_colacion', 0)
        movilizacion = emp.get('asignacion_movilizacion', 0)
        
        def safe_upper(v, default=""):
            if v is None: return default.upper()
            return str(v).upper()

        if emp.get('gratificacion_legal', True):
            clausula_gratificacion = "EL EMPLEADOR PAGARÁ AL TRABAJADOR LA GRATIFICACIÓN LEGAL EN LA MODALIDAD DEL ARTÍCULO 47 DEL CÓDIGO DEL TRABAJO, ESTO ES, EL 25% DE LO DEVENGADO POR CONCEPTO DE REMUNERACIONES MENSUALES, CON TOPE DE 4,75 INGRESOS MÍNIMOS MENSUALES."
        else:
            clausula_gratificacion = "LA GRATIFICACIÓN NO ESTÁ PACTADA BAJO LA MODALIDAD LEGAL DEL ARTÍCULO 47."

        tipo_con = emp.get('tipo_contrato', 'indefinido')
        if tipo_con == 'plazo_fijo' and emp.get('fecha_termino'):
            duracion_texto = f"EL PRESENTE CONTRATO TENDRÁ UNA DURACIÓN HASTA EL DÍA {format_date_cl(emp.get('fecha_termino')).upper()}."
        else:
            duracion_texto = "EL PRESENTE CONTRATO TENDRÁ UNA DURACIÓN INDEFINIDA."

        fecha_legal = format_date_cl(date.today())

        # 4. Contexto para la plantilla
        context = {
            'EMPRESA_NOMBRE': safe_upper(org.get('nombre'), 'EMPRESA NO REGISTRADA'),
            'EMPRESA_RUT': clean_rut(org.get('rut_empresa', '')),
            'EMPRESA_DIRECCION': safe_upper(org.get('direccion', 'DIRECCION NO REGISTRADA')),
            'EMPRESA_GIRO': safe_upper(org.get('giro'), 'ACTIVIDADES DE CONTABILIDAD'),
            'CIUDAD': safe_upper(org.get('comuna'), 'PUNTA ARENAS'),
            
            'REP_LEGAL_NOMBRE': safe_upper(settings.get('rep_legal_nombre') or settings.get('rep_nombre'), ''),
            'REP_LEGAL_RUT': clean_rut(settings.get('rep_legal_rut') or settings.get('rep_rut', '')),
            'REP_LEGAL_CARGO': safe_upper(settings.get('rep_legal_cargo') or settings.get('rep_cargo'), 'GERENTE GENERAL'),
            
            'EMPLEADO_NOMBRE': f"{emp.get('nombres', '')} {emp.get('apellido_paterno', '')} {emp.get('apellido_materno', '')}".strip().upper(),
            'EMPLEADO_RUT': clean_rut(emp.get('rut', '')),
            'EMPLEADO_NACIONALIDAD': safe_upper(emp.get('nacionalidad'), 'CHILENA'),
            'EMPLEADO_ESTADO_CIVIL': safe_upper(emp.get('estado_civil'), 'SOLTERO(A)'),
            'EMPLEADO_FECHA_NAC': format_date_cl(emp.get('birth_date')),
            'EMPLEADO_DIRECCION': safe_upper(emp.get('address') or emp.get('direccion'), 'DOMICILIO CONOCIDO'),
            'EMPLEADO_COMUNA': safe_upper(emp.get('city'), 'PUNTA ARENAS'),
            
            'FECHA_INGRESO': format_date_cl(emp.get('fecha_ingreso')),
            'DURACION_TEXTO': duracion_texto.upper(),
            'CARGO': safe_upper(emp.get('cargo'), 'TRABAJADOR'),
            'SUELDO_BASE': format_clp(sueldo_base),
            'SUELDO_PALABRAS': to_words(sueldo_base),
            'GRATIFICACION_CLAUSULA': clausula_gratificacion.upper(),
            
            'COLACION_MONTO': format_clp(colacion),
            'COLACION_PALABRAS': to_words(colacion),
            'MOVILIZACION_MONTO': format_clp(movilizacion),
            'MOVILIZACION_PALABRAS': to_words(movilizacion),
            
            'HORAS_SEMANALES': emp.get('horas_semanales', 42),
            'HORARIO': safe_upper(emp.get('horario_detalle'), 'JORNADA ORDINARIA LEGAL'),
            'PREVISION_SALUD': safe_upper(emp.get('prevision_salud'), 'FONASA'),
            'AFP': safe_upper(emp.get('afp'), 'HABITAT'),
            'FECHA_ACTUAL': fecha_legal.upper()
        }

        # 5. Cargar y renderizar plantilla
        template_name = "contrato_base.docx" if type == "contrato" else "anexo_base.docx"
        template_path = os.path.join(TEMPLATES_DIR, template_name)
        
        if not os.path.exists(template_path):
             raise HTTPException(status_code=500, detail=f"Plantilla {template_name} no encontrada.")

        doc = DocxTemplate(template_path)
        doc.render(context)
        
        # --- INCORPORAR FIRMA ---
        if sig_raw:
            try:
                if "," in sig_raw:
                    sig_raw = sig_raw.split(",")[1]
                sig_data = base64.b64decode(sig_raw)
                sig_stream = io.BytesIO(sig_data)
                
                doc.docx.add_paragraph("\n" * 2)
                p_sig = doc.docx.add_paragraph()
                p_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_sig.add_run().add_picture(sig_stream, width=Inches(1.8))
                
                # QR de verificación
                verify_url = f"https://contapymepuq.cl/verify/{employee_id[:12]}"
                qr_api_url = f"https://api.qrserver.com/v1/create-qr-code/?size=150x150&data={verify_url}"
                qr_res = requests.get(qr_api_url)
                if qr_res.status_code == 200:
                    p_sig.add_run("   ")
                    p_sig.add_run().add_picture(io.BytesIO(qr_res.content), width=Inches(0.8))
            except Exception:
                pass

        mem_file = BytesIO()
        doc.save(mem_file)
        mem_file.seek(0)
        
        nombre_archivo = f"{type.capitalize()}_{emp.get('apellido_paterno')}.docx"

        return StreamingResponse(
            iter([mem_file.getvalue()]), 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
            headers={'Content-Disposition': f'attachment; filename="{nombre_archivo}"'}
        )

    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error en motor: {str(e)}")

@router.get("/generate-annex")
async def generate_annex(
    mod_id: str,
    current_user: dict = Depends(verify_token)
):
    """
    Generador de Anexos Profesional.
    """
    db = get_supabase()
    
    try:
        mod_res = db.table("contract_modifications").select("*").eq("id", mod_id).single().execute()
        mod = mod_res.data
        if not mod: raise HTTPException(status_code=404, detail="Modificación no encontrada")
        
        employee_id = mod.get('employee_id')
        org_id = mod.get('organization_id')
        if org_id:
            await verify_org_role(org_id, auth=current_user)

        emp_res = db.table("employees").select("*").eq("id", employee_id).single().execute()
        emp = emp_res.data or {}

        org_res = db.table("organizations").select("*").eq("id", org_id).single().execute()
        org = org_res.data or {}

        settings_res = db.table("organization_payroll_settings").select("*").eq("organization_id", org_id).maybe_single().execute()
        settings = settings_res.data or {}

        # Lógica de Redacción
        changes = mod.get('changes', {})
        old_values = mod.get('old_values', {})
        clausulas = []
        
        for key, new_val in changes.items():
            old_val = old_values.get(key, "A DEFINIR")
            if key == 'sueldo_base':
                txt = f"EL SUELDO BASE, EL CUAL ASCENDÍA A {format_clp(old_val)} ({to_words(old_val)}), SE MODIFICA A {format_clp(new_val)} ({to_words(new_val)})."
            elif key == 'horas_semanales':
                txt = f"LA JORNADA PASA DE {old_val} A {new_val} HORAS SEMANALES."
            elif key == 'cargo':
                txt = f"EL CARGO PASA DE '{str(old_val).upper()}' A '{str(new_val).upper()}'."
            else:
                txt = f"SE MODIFICA '{key.upper()}', PASANDO DE '{old_val}' A '{new_val}'."
            clausulas.append(txt.upper())

        context = {
            'EMPRESA_NOMBRE': str(org.get('nombre', '')).upper(),
            'EMPRESA_RUT': clean_rut(org.get('rut_empresa', '')),
            'EMPRESA_DIRECCION': str(org.get('direccion', '')).upper(),
            'CIUDAD': str(org.get('comuna', 'PUNTA ARENAS')).upper(),
            'REP_LEGAL_NOMBRE': str(settings.get('rep_legal_nombre', '_________________')).upper(),
            'REP_LEGAL_RUT': clean_rut(settings.get('rep_legal_rut', '')),
            'EMPLEADO_NOMBRE': f"{emp.get('nombres')} {emp.get('apellido_paterno')}".upper(),
            'EMPLEADO_RUT': clean_rut(emp.get('rut', '')),
            'FECHA_ANEXO': format_date_cl(mod.get('effective_date')),
            'CLAUSULA_MODIFICACION': " ".join(clausulas),
            'FECHA_ACTUAL': format_date_cl(date.today()).upper()
        }

        template_path = os.path.join(TEMPLATES_DIR, "anexo_base.docx")
        doc = DocxTemplate(template_path)
        doc.render(context)
        mem_file = BytesIO()
        doc.save(mem_file)
        mem_file.seek(0)
        
        return StreamingResponse(
            iter([mem_file.getvalue()]), 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
            headers={'Content-Disposition': f'attachment; filename="Anexo_{emp.get("apellido_paterno")}.docx"'}
        )
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error Render: {str(e)}")
