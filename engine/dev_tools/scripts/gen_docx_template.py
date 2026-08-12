from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_contract_template():
    doc = Document()
    
    # Estilo general
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Título
    title = doc.add_paragraph('{{TITULO_DOCUMENTO}}')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)

    # Introducción
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run('En la ciudad de ')
    p.add_run('{{CIUDAD}}').bold = True
    p.add_run(', a ')
    p.add_run('{{FECHA_ACTUAL}}').bold = True
    p.add_run(', entre ')
    p.add_run('{{EMPRESA_NOMBRE}}').bold = True
    p.add_run(', RUT ')
    p.add_run('{{EMPRESA_RUT}}').bold = True
    p.add_run(', representada por don (ña) ')
    p.add_run('{{REP_LEGAL_NOMBRE}}').bold = True
    p.add_run(', RUT ')
    p.add_run('{{REP_LEGAL_RUT}}').bold = True
    p.add_run(', en su calidad de ')
    p.add_run('{{REP_LEGAL_CARGO}}').bold = True
    p.add_run(', ambos domiciliados para estos efectos en ')
    p.add_run('{{EMPRESA_DIRECCION}}').bold = True
    p.add_run(', en adelante "el Empleador"; y don (ña) ')
    p.add_run('{{EMPLEADO_NOMBRE}}').bold = True
    p.add_run(', RUT ')
    p.add_run('{{EMPLEADO_RUT}}').bold = True
    p.add_run(', domiciliado en ')
    p.add_run('{{EMPLEADO_DIRECCION}}').bold = True
    p.add_run(', en adelante "el Trabajador", se ha acordado el siguiente contrato de trabajo:')

    # Cláusulas
    clauses = [
        ('PRIMERO: Naturaleza de los servicios y polifuncionalidad.', 
         'El Trabajador se obliga a desempeñar las funciones correspondientes al cargo de {{CARGO}}, {{FUNCIONES_DETALLE}} ejecutando con la debida diligencia y lealtad las labores encomendadas por la jefatura. Asimismo, de conformidad a lo dispuesto en el inciso 3º del artículo 10 del Código del Trabajo, el Trabajador podrá realizar otras labores análogas y/o complementarias a las de su especialidad que no importen menoscabo moral o material.'),
        
        ('SEGUNDO: Lugar de desempeño.', 
         'Los servicios se prestarán en el establecimiento del Empleador ubicado en {{EMPRESA_DIRECCION}}, comuna de {{EMPLEADO_COMUNA}}, ciudad de {{CIUDAD}}, sin perjuicio de la facultad del Empleador contemplada en el artículo 12 del Código del Trabajo para alterar la naturaleza de los servicios o el sitio o recinto en que ellos deban prestarse.'),
        
        ('TERCERO: Jornada de trabajo y descanso (Ley 21.561).', 
         'La jornada ordinaria de trabajo será de {{HORAS_SEMANALES}} horas semanales, ajustada a la gradualidad establecida en la Ley Nº 21.561 de 40 Horas. Dicha jornada se distribuirá de la siguiente manera: {{HORARIO}}. El tiempo destinado a colación no se considerará trabajado ni imputable a la jornada. El control de asistencia se registrará mediante el sistema oficial de la empresa.'),
        
        ('CUARTO: Remuneraciones y modalidad de pago.', 
         'El Empleador pagará al Trabajador una remuneración mensual compuesta por:\n'
         'a) Sueldo Base mensual de {{SUELDO_BASE}} ({{SUELDO_PALABRAS}}).\n'
         'b) Gratificación Legal equivalente al 25% de la remuneración mensual con el tope de 4,75 Ingresos Mínimos Mensuales (Modalidad Art. 50 del Código del Trabajo).\n'
         'c) Asignación de Colación no imponible por {{COLACION_MONTO}} ({{COLACION_PALABRAS}}).\n'
         'd) Asignación de Movilización no imponible por {{MOVILIZACION_MONTO}} ({{MOVILIZACION_PALABRAS}}).\n'
         'Las remuneraciones se pagarán por período vencido el último día hábil de cada mes mediante transferencia electrónica a la cuenta bancaria informada por el Trabajador.'),
        
        ('QUINTO: Vigencia y duración.', 
         'El presente contrato tendrá vigencia a contar del {{FECHA_INGRESO}} y su duración será de carácter {{DURACION_TEXTO}}.'),
        
        ('SEXTO: Confidencialidad, propiedad intelectual y secreto industrial.', 
         'El Trabajador se obliga a guardar absoluta reserva y confidencialidad respecto de toda la información comercial, técnica, código fuente, listas de clientes y secretos de la empresa a los que tenga acceso con motivo de sus servicios. Asimismo, todas las creaciones, desarrollos de software y marcas generadas pertenecerán exclusivamente al Empleador conforme a la Ley Nº 17.336.'),
        
        ('SÉPTIMO: Obligación de higiene, seguridad y prevención (Ley Karin 21.643).', 
         'El Trabajador se compromete a dar estricto cumplimiento al Reglamento Interno de Orden, Higiene y Seguridad (RIOHS) de la empresa, así como a las disposiciones del Protocolo de Prevención de Acoso Laboral, Sexual y Violencia en el Trabajo establecido por la Ley Nº 21.643 (Ley Karin).'),
        
        ('OCTAVO: Protección de datos personales.', 
         'El Trabajador autoriza expresamente al Empleador para el tratamiento y almacenamiento de sus datos personales conforme a la Ley Nº 19.628, con el exclusivo objeto de dar cumplimiento a las obligaciones laborales, previsionales y tributarias.'),
        
        ('NOVENO: Domicilio y jurisdicción.', 
         'Para todos los efectos legales derivados del presente contrato, las partes fijan su domicilio en la ciudad de {{CIUDAD}} y se someten a la competencia de sus Tribunales de Letras del Trabajo.'),
        
        ('DÉCIMO: Ejemplares y Registro Laboral Electrónico (LRE).', 
         'El presente contrato se firma en dos ejemplares de idéntico tenor y fecha, quedando uno en poder de cada parte. De conformidad al artículo 9º del Código del Trabajo, el Empleador procederá al registro electrónico del presente instrumento en la plataforma de la Dirección del Trabajo (LRE).')
    ]

    for title_text, body_text in clauses:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = cp.add_run(title_text)
        run.bold = True
        doc.add_paragraph(body_text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Firmas
    doc.add_paragraph('\n\n\n')
    table = doc.add_table(rows=1, cols=2)
    table.width = Inches(6)
    
    cells = table.rows[0].cells
    
    # Firma Empleador
    p1 = cells[0].paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.add_run('__________________________\n').bold = True
    p1.add_run('FIRMA EMPLEADOR\n').bold = True
    p1.add_run('{{REP_LEGAL_NOMBRE}}\n')
    p1.add_run('RUT: {{REP_LEGAL_RUT}}\n')
    p1.add_run('p.p. {{EMPRESA_NOMBRE}}').italic = True

    # Firma Trabajador
    p2 = cells[1].paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run('__________________________\n').bold = True
    p2.add_run('FIRMA TRABAJADOR\n').bold = True
    p2.add_run('{{EMPLEADO_NOMBRE}}\n')
    p2.add_run('RUT: {{EMPLEADO_RUT}}')

    # Pie de Página (SELLO DE CALIDAD)
    section = doc.sections[0]
    footer = section.footer
    pf = footer.paragraphs[0]
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = pf.add_run('{{EMPRESA_NOMBRE}} - RUT: {{EMPRESA_RUT}} | Respaldado y Auditado Digitalmente por Contapymepuq\nEste instrumento forma parte del Kardex Laboral Electrónico de la empresa.')
    run_f.font.size = Pt(8)
    from docx.shared import RGBColor
    run_f.font.color.rgb = RGBColor(100, 100, 100) # Gris Profesional

    # Guardar
    templates_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'templates')
    if not os.path.exists(templates_path):
        os.makedirs(templates_path)
    
    file_path = os.path.join(templates_path, 'contrato_base.docx')
    doc.save(file_path)
    print(f"[OK] Plantilla creada en: {file_path}")

def create_annex_template():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Título
    title = doc.add_paragraph('ANEXO AL CONTRATO DE TRABAJO')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)

    # Introducción
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run('En la ciudad de ')
    p.add_run('{{EMPRESA_COMUNA}}').bold = True
    p.add_run(', a ')
    p.add_run('{{FECHA_ACTUAL}}').bold = True
    p.add_run(', entre ')
    p.add_run('{{EMPRESA_NOMBRE}}').bold = True
    p.add_run(', RUT ')
    p.add_run('{{EMPRESA_RUT}}').bold = True
    p.add_run(', representada por ')
    p.add_run('{{REP_LEGAL_NOMBRE}}').bold = True
    p.add_run(', RUT ')
    p.add_run('{{REP_LEGAL_RUT}}').bold = True
    p.add_run(', en su calidad de ')
    p.add_run('{{REP_LEGAL_CARGO}}').bold = True
    p.add_run(', domiciliados en ')
    p.add_run('{{EMPRESA_DIRECCION}}').bold = True
    p.add_run(', en adelante "el Empleador"; y don (ña) ')
    p.add_run('{{EMPLEADO_NOMBRE}}').bold = True
    p.add_run(', RUT ')
    p.add_run('{{EMPLEADO_RUT}}').bold = True
    p.add_run(', de nacionalidad ')
    p.add_run('{{EMPLEADO_NACIONALIDAD}}').bold = True
    p.add_run(', domiciliado en ')
    p.add_run('{{EMPLEADO_DIRECCION}}').bold = True
    p.add_run(', en adelante "el Trabajador", se ha acordado el siguiente anexo de contrato de trabajo:')

    # Cláusulas Dinámicas
    doc.add_paragraph('\nANTECEDENTE:').runs[0].bold = True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.add_run('Con motivo de ')
    p2.add_run('{{MOTIVO_ANEXO}}').bold = True
    p2.add_run(', las partes han resuelto modificar las condiciones del contrato de trabajo vigente.')

    doc.add_paragraph('\nPRIMERO: MODIFICACIÓN').runs[0].bold = True
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p3.add_run('{{CLAUSULA_MODIFICACION}}')

    doc.add_paragraph('\nSEGUNDO: VIGENCIA').runs[0].bold = True
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p4.add_run('La presente modificación entrará en vigencia a contar del día ')
    p4.add_run('{{FECHA_ANEXO}}').bold = True
    p4.add_run('.')

    doc.add_paragraph('\nTERCERO: MANTENCIÓN DE CONDICIONES').runs[0].bold = True
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p5.add_run('{{CL_VIGENCIA}}')

    # Firmas
    doc.add_paragraph('\n\n\n')
    table = doc.add_table(rows=1, cols=2)
    table.width = Inches(6)
    cells = table.rows[0].cells
    
    p1 = cells[0].paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.add_run('__________________________\n').bold = True
    p1.add_run('FIRMA EMPLEADOR\n').bold = True
    p1.add_run('{{REP_LEGAL_NOMBRE}}\n')
    p1.add_run('RUT: {{REP_LEGAL_RUT}}\n')
    p1.add_run('p.p. {{EMPRESA_NOMBRE}}').italic = True

    p2 = cells[1].paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run('__________________________\n').bold = True
    p2.add_run('FIRMA TRABAJADOR\n').bold = True
    p2.add_run('{{EMPLEADO_NOMBRE}}\n')
    p2.add_run('RUT: {{EMPLEADO_RUT}}')

    # Pie de Página (SELLO DE CALIDAD ANEXO)
    section = doc.sections[0]
    footer = section.footer
    pf = footer.paragraphs[0]
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = pf.add_run('{{EMPRESA_NOMBRE}} - RUT: {{EMPRESA_RUT}} | Respaldado y Auditado Digitalmente por Contapymepuq\nEste instrumento forma parte del Kardex Laboral Electrónico de la empresa.')
    run_f.font.size = Pt(8)
    from docx.shared import RGBColor
    run_f.font.color.rgb = RGBColor(100, 100, 100) # Gris Profesional

    # Guardar en la carpeta correcta
    templates_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'templates')
    if not os.path.exists(templates_path):
        os.makedirs(templates_path)
    
    file_path = os.path.join(templates_path, 'anexo_base.docx')
    doc.save(file_path)
    print(f"[OK] Plantilla Anexo creada en: {file_path}")

if __name__ == "__main__":
    create_contract_template()
    create_annex_template()
