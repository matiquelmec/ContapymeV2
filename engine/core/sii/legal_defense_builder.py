"""
legal_defense_builder.py — Generador de Escritos Legales Formales para el SII
Produce documentos Word (.docx) dirigidos al Director Regional del SII de Magallanes.
Respeta los estándares de formato formal: Times New Roman 11pt, márgenes 3.5cm/2.5cm,
citas del D.L. 825 (Ley del IVA) y D.L. 830 (Código Tributario).
"""

import io
import datetime
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

MESES = {
    1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
    5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
    9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre"
}

def format_monto_clp(v: Any) -> str:
    try:
        return f"${int(round(float(v))):,}".replace(",", ".")
    except (TypeError, ValueError):
        return "$0"

class SIILegalDefenseBuilder:
    @staticmethod
    def build_docx(data: Dict[str, Any]) -> io.BytesIO:
        """
        Construye el documento DOCX oficial.
        data keys:
        - doc_type: 'boletas_vs_facturas' | 'citacion_art_63' | 'rectificatoria_f29' | 'condonacion_multas'
        - razon_social: str
        - rut_empresa: str
        - rep_legal: str
        - rep_rut: str
        - domicilio: str (Punta Arenas)
        - giro: str
        - periodos: List[str] (ej: ['2025-01', '2025-02'])
        - iva_declarado: float
        - numero_citacion: Optional[str]
        - argumentos_adicionales: Optional[str]
        """
        doc = Document()
        doc_type = data.get("doc_type", "boletas_vs_facturas")
        
        # 1. Configuración de Márgenes Reglamentarios
        for sec in doc.sections:
            sec.top_margin    = Cm(3.0)
            sec.bottom_margin = Cm(2.5)
            sec.left_margin   = Cm(3.5)
            sec.right_margin  = Cm(2.5)

        def add_p(text: str = "", bold: bool = False, italic: bool = False, size: int = 11,
                  align = WD_ALIGN_PARAGRAPH.LEFT, space_before: int = 0, space_after: int = 8,
                  indent_cm: float = 0):
            p = doc.add_paragraph()
            p.alignment = align
            p.paragraph_format.space_before = Pt(space_before)
            p.paragraph_format.space_after  = Pt(space_after)
            p.paragraph_format.line_spacing = 1.15
            if indent_cm > 0:
                p.paragraph_format.left_indent = Cm(indent_cm)
            if text:
                r = p.add_run(text)
                r.bold = bold
                r.italic = italic
                r.font.size = Pt(size)
                r.font.name = "Times New Roman"
            return p

        hoy = datetime.date.today()
        fecha_str = f"Punta Arenas, {hoy.day} de {MESES[hoy.month]} de {hoy.year}"
        
        razon_social = data.get("razon_social", "EMPRESA CONTRIBUYENTE SPA").upper()
        rut_empresa = data.get("rut_empresa", "76.000.000-0")
        rep_legal = data.get("rep_legal", "REPRESENTANTE LEGAL").upper()
        rep_rut = data.get("rep_rut", "15.000.000-0")
        domicilio = data.get("domicilio", "Punta Arenas, Región de Magallanes")
        giro = data.get("giro", "Servicios Integrales")
        periodos = data.get("periodos", [])
        periodos_str = ", ".join(periodos) if periodos else "los períodos fiscalizados"
        iva_total_str = format_monto_clp(data.get("iva_declarado", 0))

        # ── ENCABEZADO OFICIAL ──
        add_p(fecha_str, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=18)
        add_p("Señor", space_after=2)
        add_p("Director Regional", space_after=2)
        add_p("Servicio de Impuestos Internos", space_after=2)
        add_p("XII Región de Magallanes y de la Antártica Chilena", space_after=2)
        add_p("Unidad Punta Arenas", space_after=18)
        add_p("Presente", bold=True, space_after=20)

        # ── IDENTIFICACIÓN Y COMPARECENCIA ──
        texto_comparecencia = (
            f"{razon_social}, Rol Único Tributario N° {rut_empresa}, del giro de su denominación, "
            f"representada legalmente por don(ña) {rep_legal}, cédula nacional de identidad N° {rep_rut}, "
            f"ambos domiciliados para estos efectos en {domicilio}, Punta Arenas, ante usted respetuosamente "
            f"expone y presenta:"
        )
        add_p(texto_comparecencia, space_after=16)

        # ── CONTENIDO SEGÚN TIPO DE CAUSA ──
        if doc_type == "boletas_vs_facturas":
            add_p("I.  ANTECEDENTES Y REQUERIMIENTO DE FISCALIZACIÓN", bold=True, space_after=8)
            add_p(
                f"Con motivo de actuaciones de fiscalización practicadas por funcionarios de esta Dirección Regional "
                f"respecto de las operaciones comerciales de mi representada, se ha formulado observación o requerimiento "
                f"relativo a la presunta falta de emisión de facturas de venta en los períodos {periodos_str}.",
                space_after=10
            )
            add_p(
                "Al respecto, vengo en formular la presente presentación voluntaria de descargos con el objeto de "
                "acreditar de manera fehaciente que dicho requerimiento no es aplicable a la naturaleza de las operaciones "
                "del contribuyente y que el interés fiscal se encuentra íntegramente resguardado.",
                space_after=16
            )

            add_p("II.  DEL GIRO Y NATURALEZA DE LAS VENTAS A CONSUMIDORES FINALES", bold=True, space_after=8)
            add_p(
                f"2.1  Mi representada tiene como actividad principal el giro de '{giro}', en cuya virtud la totalidad de "
                f"las ventas y servicios son prestados de manera directa y exclusiva a consumidores finales.",
                space_after=10
            )
            add_p(
                "2.2  De conformidad con lo dispuesto en el artículo 53 del D.L. N° 825 sobre Impuesto a las Ventas y Servicios, "
                "en concordancia con el artículo 35 de su Reglamento, los contribuyentes que efectúen transferencias de bienes "
                "o presten servicios a consumidores finales están legalmente obligados a emitir boletas de ventas y servicios, "
                "siendo la factura un instrumento exigible únicamente respecto de operaciones entre contribuyentes de IVA.",
                space_after=10
            )
            add_p(
                "2.3  En consecuencia, exigir la emisión de facturas respecto de ventas efectuadas a personas naturales en calidad "
                "de consumidores finales carece de fundamento legal y contraviene la normativa tributaria expresa.",
                space_after=16
            )

            add_p("III.  DECLARACIÓN Y PAGO ÍNTEGRO DEL IVA: AUSENCIA DE PERJUICIO FISCAL", bold=True, space_after=8)
            add_p(
                f"3.1  Consta en los registros del propio Servicio que la totalidad de los ingresos percibidos durante "
                f"{periodos_str} han sido íntegramente declarados en los respectivos Formularios 29 presentados mensualmente.",
                space_after=10
            )
            add_p(
                f"3.2  El Impuesto al Valor Agregado (Débito Fiscal) correspondiente a las citadas operaciones, ascendente a un total de "
                f"{iva_total_str}, fue debidamente determinado y enterado en arcas fiscales en tiempo y forma, no habiendo existido "
                f"evasión ni detrimento fiscal alguno.",
                space_after=16
            )

        elif doc_type == "citacion_art_63":
            num_cit = data.get("numero_citacion", "S/N")
            add_p(f"I.  CUMPLIMIENTO A CITACIÓN N° {num_cit} (ART. 63 CÓDIGO TRIBUTARIO)", bold=True, space_after=8)
            add_p(
                f"Dentro del plazo legal conferido en la Citación N° {num_cit}, vengo en dar oportuna y cabal respuesta a cada uno "
                f"de los reparos formulados por el fiscalizador actuante respecto de los períodos {periodos_str}.",
                space_after=10
            )
            add_p("II.  FUNDAMENTACIÓN DE HECHO Y DE DERECHO", bold=True, space_after=8)
            add_p(
                f"Las partidas observadas se encuentran debidamente respaldadas en la contabilidad fidedigna de la empresa, "
                f"en el Registro de Compras y Ventas (RCV) electrónico y en las declaraciones mensuales F29, demostrando la correlación "
                f"exacta entre los débitos y créditos declarados y el giro efectivo de la sociedad.",
                space_after=16
            )

        elif doc_type == "rectificatoria_f29":
            add_p("I.  SOLICITUD DE RECTIFICATORIA VOLUNTARIA (ART. 127 CÓDIGO TRIBUTARIO)", bold=True, space_after=8)
            add_p(
                f"Al amparo de lo prescrito en el artículo 127 del D.L. N° 830 (Código Tributario), comparezco solicitando "
                f"autorización y validación formal de la rectificatoria de los Formularios 29 correspondientes a {periodos_str}.",
                space_after=10
            )
            add_p(
                "La presente rectificación obedece a un mero error de hecho involuntario en la imputación de códigos, sin que "
                "ello haya alterado la base imponible real ni la determinación del impuesto adeudado.",
                space_after=16
            )

        elif doc_type == "condonacion_multas":
            add_p("I.  SOLICITUD DE CONDONACIÓN DE INTERESES Y MULTAS (CIRCULAR N° 50 SII)", bold=True, space_after=8)
            add_p(
                "En virtud de las atribuciones otorgadas al Director Regional en el artículo 6° letra B N° 4 del Código Tributario, "
                "y de conformidad con los criterios e instrucciones contenidos en la Circular N° 50 del Servicio de Impuestos Internos, "
                "vengo en solicitar la condonación del máximo porcentaje legal de los intereses penales y multas.",
                space_after=10
            )
            add_p(
                "Funda esta petición el historial intachable de cumplimiento de mi representada, la corrección espontánea y voluntaria "
                "de cualquier observación y la evidente buena fe exenta de toda maniobra de elusión o fraude tributario.",
                space_after=16
            )

        # ── ARGUMENTOS ADICIONALES (SI LOS HAY) ──
        if data.get("argumentos_adicionales"):
            add_p("DEL DETALLE ESPECÍFICO DE ANTECEDENTES", bold=True, space_after=8)
            add_p(data["argumentos_adicionales"], space_after=16)

        # ── PETICIÓN FINAL (PETITUM FORMAL) ──
        add_p("POR TANTO,", bold=True, space_after=8)
        add_p(
            f"En mérito de lo expuesto, disposiciones legales citadas y antecedentes que se acompañan, "
            f"a usted respetuosamente pido:\n\n"
            f"Se sirva tener por presentado este escrito de descargos y aclaración en favor de {razon_social}, "
            f"tener por acreditada la legalidad de las operaciones comerciales de los períodos {periodos_str}, "
            f"constatar la inexistencia de perjuicio fiscal y ordenar el levantamiento definitivo de toda observación o reparo "
            f"recaído sobre mi representada, emitiendo la resolución que en derecho corresponda.",
            space_after=28
        )

        # ── FIRMA ──
        add_p("____________________________________________", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
        add_p(f"{rep_legal}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
        add_p(f"RUT: {rep_rut}", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
        add_p(f"Representante Legal de {razon_social}", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

        # ── GUARDAR EN MEMORIA ──
        target_stream = io.BytesIO()
        doc.save(target_stream)
        target_stream.seek(0)
        return target_stream
