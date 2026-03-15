from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_contract_template():
    doc = Document()
    
    # Estilo general
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Título
    title = doc.add_paragraph('CONTRATO DE TRABAJO')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)

    # Introducción
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run('En la ciudad de ')
    p.add_run('{{CIUDAD}}').bold = True
    p.add_run(', a ')
    p.add_run('{{FECHA_ACTUAL}}').bold = True
    p.add_run(', entre ')
    p.add_run('{{EMPRESA_NOMBRE}}').bold = True
    p.add_run(', RUT ')
    p.add_run('{{EMPRESA_RUT}}').bold = True
    p.add_run(', representada por don (ña) ')
    p.add_run('{{REP_LEGAL_NOMBRE}}').bold = True
    p.add_run(', RUT ')
    p.add_run('{{REP_LEGAL_RUT}}').bold = True
    p.add_run(', en su calidad de ')
    p.add_run('{{REP_LEGAL_CARGO}}').bold = True
    p.add_run(', ambos domiciliados para estos efectos en ')
    p.add_run('{{EMPRESA_DIRECCION}}').bold = True
    p.add_run(', en adelante "el Empleador"; y don (ña) ')
    p.add_run('{{EMPLEADO_NOMBRE}}').bold = True
    p.add_run(', RUT ')
    p.add_run('{{EMPLEADO_RUT}}').bold = True
    p.add_run(', domiciliado en ')
    p.add_run('{{EMPLEADO_DIRECCION}}').bold = True
    p.add_run(', en adelante "el Trabajador", se ha acordado el siguiente contrato de trabajo:')

    # Cláusulas
    clauses = [
        ('PRIMERO: Naturaleza de los servicios.', 
         'El trabajador se obliga a desempeñar el cargo de {{CARGO}} y todas las funciones inherentes a dicha posición, así como aquellas que el empleador le encomiende de acuerdo a su especialidad y conocimientos.'),
        
        ('SEGUNDO: Lugar de desempeño.', 
         'Los servicios se prestarán en las dependencias de la empresa ubicadas en {{EMPRESA_DIRECCION}}, ciudad de {{CIUDAD}}, sin perjuicio de lo dispuesto en el artículo 12 del Código del Trabajo.'),
        
        ('TERCERO: Jornada de trabajo.', 
         'La jornada de trabajo será de 44 horas semanales, distribuidas de lunes a viernes en horario general de la empresa.'),
        
        ('CUARTO: Remuneraciones.', 
         'El empleador pagará al trabajador una remuneración mensual de {{SUELDO_BASE}} ({{SUELDO_PALABRAS}}), la que será liquidada y pagada el último día hábil de cada mes.'),
        
        ('QUINTO: Vigencia del contrato.', 
         'El presente contrato tendrá una vigencia a contar del {{FECHA_INGRESO}} y su duración será de carácter INDEFINIDO.')
    ]

    for title_text, body_text in clauses:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = cp.add_run(title_text)
        run.bold = True
        doc.add_paragraph(body_text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Firmas
    doc.add_paragraph('\n\n\n')
    table = doc.add_table(rows=1, cols=2)
    table.width = Inches(6)
    
    cells = table.rows[0].cells
    
    # Firma Empleador
    p1 = cells[0].paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.add_run('__________________________\n').bold = True
    p1.add_run('{{EMPRESA_NOMBRE}}\n').bold = True
    p1.add_run('RUT: {{EMPRESA_RUT}}\n')
    p1.add_run('EMPLEADOR')

    # Firma Trabajador
    p2 = cells[1].paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run('__________________________\n').bold = True
    p2.add_run('{{EMPLEADO_NOMBRE}}\n').bold = True
    p2.add_run('RUT: {{EMPLEADO_RUT}}\n')
    p2.add_run('TRABAJADOR')

    # Guardar
    templates_path = os.path.join(os.getcwd(), 'templates')
    if not os.path.exists(templates_path):
        os.makedirs(templates_path)
    
    file_path = os.path.join(templates_path, 'contrato_base.docx')
    doc.save(file_path)
    print(f"✅ Plantilla creada en: {file_path}")

if __name__ == "__main__":
    create_contract_template()
